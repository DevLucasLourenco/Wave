import time
import docx
from copy import deepcopy

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from service.data.archive import Archive


def deltaTime(method):
    def wrapper(self, *args, **kwargs):
        initialTime = time.time()
        mtd = method(self, *args, **kwargs)
        self._timeToGenerate = time.time() - initialTime
        return mtd
    return wrapper


class Builder: # e se mdar pra Builder?
    
    def __init__(self, archive:Archive, baseDocx) -> None:
        self.__archive = archive
        self._timeToGenerate:int
        self.__baseDocx = docx.Document(baseDocx)
        self.__allKeys :str = list(self.__archive.getMetaData()['columns'].values())[0]
        self.__firstKey:str = self.__allKeys[0]
    
    
    @deltaTime
    def generate(self):
        for i, _ in enumerate(self.__archive.getData()[self.__firstKey]['data_column']):
            allRecordsFromIndex = self.__getRecordsFromSameIndex(i)
            doc = self.__replaceInfosAtDoc(allRecordsFromIndex)
            
            doc.save(f'teste{i}.docx') # remover para colocar uma função de salvar
            
            
    def __replaceInfosAtDoc(self, records):
        doc_base = deepcopy(self.__baseDocx)
        self.__paragraph(records, doc_base)
        self.__table(records, doc_base)
        return doc_base

                            
    def __paragraph(self, records:dict, doc_base:docx.Document):        
        for para in doc_base.paragraphs:
            for key, value in records.items():
                if key in para.text:
                    for run in para.runs:
                        if key in run.text:
                            primaryKey = self.__getPrimaryKeyFromKeyWDelimiter(key)
                            run.text = run.text.replace(key, str(value))
                            
                            additional_params = self.__archive.getData()[primaryKey]['additional_parameters']
                            
                            if additional_params['font']:
                                run.font.name = additional_params['font']
                                
                            if additional_params['size'] != 0:
                                run.font.size = Pt(additional_params['size'])
                                
                            if additional_params['bold']:
                                run.bold = True
                                
                            if additional_params['italic']:
                                run.italic = True

                
    
    def __table(self, records:dict, doc_base:docx.Document):
        for table in doc_base.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in records.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    primaryKey = self.__getPrimaryKeyFromKeyWDelimiter(key)
                                    additional_params = self.__archive.getData()[primaryKey]['additional_parameters']
                                    
                                    if additional_params['font']:
                                        run.font.name = additional_params['font']
                                        
                                    if additional_params['size'] != 0:
                                        run.font.size = Pt(additional_params['size'])
                                        
                                    if additional_params['bold']:
                                        run.bold = True
                                        
                                    if additional_params['italic']:
                                        run.italic = True

                                                
    def __getPrimaryKeyFromKeyWDelimiter(self, KeyWDelimiter:str):
            return KeyWDelimiter.strip(self.__archive.getDelimiter())
    
                        
    def __getRecordsFromSameIndex(self, index) -> dict: 
        d_aux=dict()
        for keyHeader in self.__allKeys:
            information = self.__archive.getData()[keyHeader]['data_column'][index]
            KeyWithDelimiter = self.__archive.getData()[keyHeader]['key_w/Delimiter']
            d_aux.update({KeyWithDelimiter:information})
            
        return d_aux
    
    
    def buildPDFs(self):
        ...
    
    
    def getTimeToGenerate(self):
        if self._timeToGenerate:
            return self._timeToGenerate
        