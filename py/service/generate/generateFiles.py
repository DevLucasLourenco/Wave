import time
import docx
from copy import deepcopy

from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from service.data.archive import Archive


def deltaTime(method):
    def wrapper(self, *args, **kwargs):
        initialTime = time.time()
        mtd = method(self, *args, **kwargs)
        self._timeToGenerate = time.time() - initialTime
        return mtd
    return wrapper


class Generate:
    
    def __init__(self, archive:Archive, baseDocx) -> None:
        self.__archive = archive
        self._timeToGenerate:int
        self.__baseDocx = docx.Document(baseDocx)
        self.__allKeys :str = list(self.__archive.getMetaData()['columns'].values())[0]
        self.__firstKey:str = self.__allKeys[0]
    
    
    @deltaTime
    def generate(self):
        for i, _ in enumerate(self.__archive.getData()[self.__firstKey]['data_column']):
            allRecordsFromIndex = self.__getRecordsFromSameIndex(i)
            print(allRecordsFromIndex)
            doc = self.__replaceInfosAtDoc(allRecordsFromIndex)
            
            doc.save('teste.docx')
            
            
            
    
    def __replaceInfosAtDoc(self, records):
        doc_base = deepcopy(self.__baseDocx)
        self.__paragraph(records, doc_base)
        self.__table(records, doc_base)
        return doc_base

                            
    def __paragraph(self, records:dict, doc_base:docx.Document):        
        for para in doc_base.paragraphs:
            for key, value in records.items():
                if key in para.text:
                    para.text = para.text.replace(key, str(value))
                    # for run in para.runs:
                    #     if self.__archive.getData()[self.__firstKey]['additional_parameters']['font']:
                    #         run.font.name = self.__archive.getData()[self.__firstKey]['additional_parameters']['font']
                            
                    #     if self.__archive.getData()[self.__firstKey]['additional_parameters']['size'] != 0:
                    #         run.font.size = Pt(self.__archive.getData()[self.__firstKey]['additional_parameters']['size'])
                            
                    #     if self.__archive.getData()[self.__firstKey]['additional_parameters']['bold']:
                    #         run.bold = True
                            
                    #     if self.__archive.getData()[self.__firstKey]['additional_parameters']['italic']:
                    #         run.italic = True
        
    
    def __table(self, records:dict, doc_base:docx.Document):
        for table in doc_base.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in records.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))
                            # for para in cell.paragraphs:
                            #     for run in para.runs:
                            #         if self.__archive.getData()[self.__firstKey]['additional_parameters']['font']:
                            #             run.font.name = self.__archive.getData()[self.__firstKey]['additional_parameters']['font']
                                        
                            #         if self.__archive.getData()[self.__firstKey]['additional_parameters']['size'] != 0:
                            #             run.font.size = Pt(self.__archive.getData()[self.__firstKey]['additional_parameters']['size'])
                                        
                            #         if self.__archive.getData()[self.__firstKey]['additional_parameters']['bold']:
                            #             run.bold = True
                                        
                            #         if self.__archive.getData()[self.__firstKey]['additional_parameters']['italic']:
                            #             run.italic = True
                                                
                        
                        
    def __getRecordsFromSameIndex(self, index) -> dict: 
        d_aux=dict()
        for keyHeader in self.__allKeys:
            information = self.__archive.getData()[keyHeader]['data_column'][index]
            KeyWithDelimiter = self.__archive.getData()[keyHeader]['key_w/Delimiter']
            
            d_aux.update({KeyWithDelimiter:information})
            
        return d_aux
    
    
    def buildPDFs(self):
        ...
    
    
    def getTimeToGenerate(self):
        if self._timeToGenerate:
            return self._timeToGenerate
        